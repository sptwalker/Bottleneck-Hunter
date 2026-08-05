"""C2 衍生品/结构化产品建模（M3）：Accumulator/Decumulator 与 MLI Booster 两类。

范围（先做最有价值的两族样本）：
1) Equity Accumulator / Decumulator（每日累积/减持，带 Guarantee、KO、Step-up shares）
2) Equity Market Linked Instrument / Booster（参与率 + KI Put + capped upside）

本模块两层能力：
- 术语抽取：从 term sheet / final terms PDF 文本抽关键条款，产出规范 dict。
- 场景收益：给定终值（及/或是否触发 KI/KO）算到期收益/交割股数，供报告提示风险。

说明：Accumulator/Decumulator 是路径依赖产品，不用 BS；先做静态/准静态场景引擎。
MLI Booster 是结构化票据，可抽象为 capped participation + down-and-in put 的到期收益函数。
"""
from __future__ import annotations

import math
import re
from dataclasses import dataclass
from datetime import datetime

# ── Black-Scholes 纯函数（D4 规格，先供 MLI / 后续标准期权）──────────────

def _cdf(x: float) -> float:
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def _pdf(x: float) -> float:
    return math.exp(-0.5 * x * x) / math.sqrt(2.0 * math.pi)


def bs_price(S, K, T, r, sigma, is_call, q=0.0):
    if S <= 0 or K <= 0:
        return 0.0
    if T <= 0:
        return max(S - K, 0.0) if is_call else max(K - S, 0.0)
    if sigma <= 0:
        fwd = S * math.exp(-q * T) - K * math.exp(-r * T)
        return max(fwd, 0.0) if is_call else max(-fwd, 0.0)
    d1 = (math.log(S / K) + (r - q + 0.5 * sigma * sigma) * T) / (sigma * math.sqrt(T))
    d2 = d1 - sigma * math.sqrt(T)
    if is_call:
        return S * math.exp(-q * T) * _cdf(d1) - K * math.exp(-r * T) * _cdf(d2)
    return K * math.exp(-r * T) * _cdf(-d2) - S * math.exp(-q * T) * _cdf(-d1)


def bs_greeks(S, K, T, r, sigma, is_call, q=0.0):
    price = bs_price(S, K, T, r, sigma, is_call, q)
    if S <= 0 or K <= 0 or T <= 0 or sigma <= 0:
        intrinsic_delta = 1.0 if (is_call and S > K) else (-1.0 if (not is_call and S < K) else 0.0)
        return {"price": price, "delta": intrinsic_delta, "gamma": 0.0, "vega": 0.0, "theta": 0.0, "rho": 0.0}
    d1 = (math.log(S / K) + (r - q + 0.5 * sigma * sigma) * T) / (sigma * math.sqrt(T))
    d2 = d1 - sigma * math.sqrt(T)
    nd1 = _pdf(d1)
    disc_q = math.exp(-q * T)
    disc_r = math.exp(-r * T)
    delta = disc_q * _cdf(d1) if is_call else disc_q * (_cdf(d1) - 1)
    gamma = disc_q * nd1 / (S * sigma * math.sqrt(T))
    vega = S * disc_q * nd1 * math.sqrt(T) / 100.0
    if is_call:
        theta = (-S * disc_q * nd1 * sigma / (2 * math.sqrt(T))
                 - r * K * disc_r * _cdf(d2)
                 + q * S * disc_q * _cdf(d1)) / 365.0
        rho = K * T * disc_r * _cdf(d2) / 100.0
    else:
        theta = (-S * disc_q * nd1 * sigma / (2 * math.sqrt(T))
                 + r * K * disc_r * _cdf(-d2)
                 - q * S * disc_q * _cdf(-d1)) / 365.0
        rho = -K * T * disc_r * _cdf(-d2) / 100.0
    return {"price": price, "delta": delta, "gamma": gamma, "vega": vega, "theta": theta, "rho": rho}


def implied_vol(price, S, K, T, r, is_call, q=0.0):
    if price <= 0 or S <= 0 or K <= 0 or T <= 0:
        return None
    lo, hi = 1e-4, 5.0
    # 越界：连高波动都定不到 → 无解
    if price < bs_price(S, K, T, r, lo, is_call, q) - 1e-9:
        return None
    if price > bs_price(S, K, T, r, hi, is_call, q) + 1e-9:
        return None
    for _ in range(80):
        mid = (lo + hi) / 2
        pm = bs_price(S, K, T, r, mid, is_call, q)
        if abs(pm - price) < 1e-6:
            return mid
        if pm > price:
            hi = mid
        else:
            lo = mid
    return (lo + hi) / 2


# ── 结构化产品规范模型 ─────────────────────────────────────────────────────

@dataclass
class DerivativeTerm:
    product_family: str        # equity_accumulator / equity_decumulator / equity_mli_booster
    underlying_symbol: str
    currency: str
    tenor_days: int
    terms: dict
    source_file: str = ""
    id: str = ""               # 落库记录 id（loader 回填；新抽取时为空）
    lot_key: str = ""          # 同标的多笔头寸判别键（strike:maturity 等）；单条留空


# ── 文本抽取 helper ───────────────────────────────────────────────────────

def _read_pdf_text(pdf_source, pages: int = 6, pdf_password: str = "") -> str:
    """读 PDF 前 N 页文本。pdf_source 可为 path(str/Path) 或 bytes；加密 PDF 可传密码。"""
    import fitz
    if isinstance(pdf_source, (bytes, bytearray)):
        doc = fitz.open(stream=pdf_source, filetype="pdf")
    else:
        doc = fitz.open(str(pdf_source))
    if doc.needs_pass:
        if not pdf_password or not doc.authenticate(pdf_password):
            raise ValueError("pdf_password_required_or_invalid")
    return "\n".join(page.get_text() for page in doc[:pages])


def _f(pat: str, text: str, group=1, flags=re.I) -> str | None:
    m = re.search(pat, text, flags)
    return m.group(group).strip() if m else None


def _ff(pat: str, text: str, group=1) -> float | None:
    s = _f(pat, text, group)
    if s is None:
        return None
    try:
        return float(s.replace(",", ""))
    except ValueError:
        return None


def _days_between(a: str, b: str) -> int:
    """兼容 Citi/野村日期格式：Jul 22, 2026 / 7 July 2026 / July 7, 2026。"""
    fmts = ("%b %d, %Y", "%d %B %Y", "%B %d, %Y", "%d %b %Y")
    def parse(s: str):
        for f in fmts:
            try:
                return datetime.strptime(s.strip(), f)
            except ValueError:
                continue
        raise ValueError(f"unsupported date format: {s}")
    da = parse(a)
    db = parse(b)
    return (db - da).days


# ── 条款抽取：Accumulator/Decumulator ────────────────────────────────────

def extract_accumulator_terms(pdf_source, pdf_password: str = "") -> DerivativeTerm:
    text = _read_pdf_text(pdf_source, pages=8, pdf_password=pdf_password)
    if "RATIONALE RECORD" in text:  # 野村 irf 精简记录（非完整条款单）走专用抽取
        return _parse_irf_record(text, pdf_source)
    # 识别 product family：正文常同时出现“Equity Accumulator / Equity Decumulator”说明语，
    # 故优先看产品标题里的 Daily ... Accumulator/Decumulator。
    fam = "equity_accumulator"
    if re.search(r"Daily(?: Securities)? Decumulator", text, re.I):
        fam = "equity_decumulator"
    elif re.search(r"Daily(?: Securities)? Accumulator", text, re.I):
        fam = "equity_accumulator"

    # Citi 样本：Bloomberg Ticker / AFP / KO / DS / St-DS / Max Nominal Shares
    if "Micron Technology Inc" in text or "Marvell Technology Inc" in text or "Alibaba Group Holding" in text:
        symbol = _f(r"Bloomberg Ticker\s*:?\s*([A-Z0-9]{1,6}\s+[A-Z]{2})", text) or ""
        symbol = symbol.split()[0]
        ccy = _f(r":\s*([A-Z]{3})\s+\d+(?:\.\d+)?\s*\(.*Initial Price\)", text) or "USD"
        trade_date = _f(r"Trade Date\s*:?\s*([A-Za-z]{3}\s+\d{1,2},\s+\d{4})", text) or ""
        termination = _f(r"Termination Date\s*:?\s*The earlier of \(a\)\s*([A-Za-z]{3}\s+\d{1,2},\s+\d{4})", text) or ""
        tenor = _days_between(trade_date, termination) if trade_date and termination else 365
        ds = _ff(r"Daily Number of Shares \(DS\)\s*:?\s*(\d+(?:\.\d+)?)", text) or 0.0
        # St-DS 缺抽时不可留 0：低于行权价时 payoff 用 days*St-DS 累股，St-DS=0 会把下行累购/亏损算成 0，
        # 恰在最危险方向静默低估风险。日累购市场惯例 step-up=2×DS → 缺失回落 2*DS(偏保守)。ponytail: D3 校准旋钮。
        stds = _ff(r"Step-up Daily Number of Shares \(St-DS\)\s*:?\s*(\d+(?:\.\d+)?)", text) or (2.0 * ds)
        max_nom = _ff(r"Maximum Number of Nominal Shares\s*:?\s*([\d,]+(?:\.\d+)?)", text) or 0.0
        afp = _ff(r":\s*USD\s*([\d,]+\.\d+)\s*\(\s*70\.75% of Initial Price\s*\)", text) or _ff(r"AFP\)?\s*:?\s*USD\s*([\d,]+\.\d+)", text) or 0.0
        initial = _ff(r"Initial Price\s*:?\s*USD\s*([\d,]+\.\d+)", text) or 0.0
        ko = _ff(r"Knock-out Price \(KO\)\s*:?\s*USD\s*([\d,]+\.\d+)", text) or 0.0
        gp = _f(r"Guaranteed Period End Date\s*:?\s*([A-Za-z]{3}\s+\d{1,2},\s+\d{4})", text) or ""
        return DerivativeTerm(
            product_family=fam, underlying_symbol=symbol, currency=ccy, tenor_days=tenor, source_file=str(pdf_source),
            terms={"initial_price": initial, "afp": afp, "knock_out_price": ko,
                   "knock_out_direction": "down_and_out" if fam == "equity_decumulator" else "up_and_out",
                   "daily_shares": ds, "step_up_daily_shares": stds, "max_nominal_shares": max_nom,
                   "guaranteed_period_end": gp, "settlement_style": "physical_spot", "net_premium": 0.0,
                   "trade_date": trade_date, "expiry_date": termination, "tenor_days": tenor},
        )

    # Nomura 样本：12 Month USD Daily Accumulator/Decumulator（BE.N / PLTR.OQ）
    symbol = _f(r"Underlying Share\s*([A-Za-z0-9 .&'/-]+)?\s*\(([A-Z0-9]{1,6})\s+[A-Z]{2}\s+Equity\)", text, group=2, flags=re.I | re.S) or ""
    if not symbol:
        symbol = (_f(r"^([A-Z0-9.]{2,10}),\s*\d+(?:\.\d+)?%\s*Strike Price", text, flags=re.I | re.M) or "").split('.')[0]
    ccy = _f(r"Settlement Currency\s*([A-Z]{3})", text) or _f(r"Underlying CCY\s*([A-Z]{3})", text) or "USD"
    trade_date = _f(r"Trade Date\s*([0-9]{1,2} [A-Za-z]+ 20\d{2})", text) or ""
    end_pat = r"Final Decumulation Date\s*([0-9]{1,2} [A-Za-z]+ 20\d{2})" if fam == "equity_decumulator" else r"Final Accumulation Date\s*([0-9]{1,2} [A-Za-z]+ 20\d{2})"
    final_date = _f(end_pat, text) or ""
    tenor = _days_between(trade_date, final_date) if trade_date and final_date else 365
    initial = _ff(r"Reference Spot Price \(USD\)\s*([\d,]+\.\d+)", text) or 0.0
    afp = _ff(r"Forward Price \(USD\)\s*([\d,]+\.\d+)", text) or 0.0
    ko = _ff(r"Knock(?:-Out)? (?:Price|Level) \(USD\)\s*([\d,]+\.\d+)", text) or 0.0
    max_nom = _ff(r"Maximum Total Shares\s*([\d,]+(?:\.\d+)?)", text) or 0.0
    ds = _ff(r"Shares per Day\s*([\d,]+(?:\.\d+)?)", text) or _ff(r"Shares per day\s*([\d,]+(?:\.\d+)?)", text) or 0.0
    gear = _ff(r"Gearing Ratio\s*([\d,]+(?:\.\d+)?)", text) or 1.0
    # Nomura 日累积/减持没有显式 Step-up shares，而是 LNBD * Gearing Ratio —— 折算为 step-up daily shares = DS*GearingRatio
    stds = ds * gear
    protected_end = _f(r"Protected\s+Period\s+End\s+Date.*?([0-9]{1,2} [A-Za-z]+ 20\d{2})", text) or ""
    return DerivativeTerm(
        product_family=fam, underlying_symbol=symbol, currency=ccy, tenor_days=tenor, source_file=str(pdf_source),
        terms={"initial_price": initial, "afp": afp, "knock_out_price": ko,
               "knock_out_direction": "down_and_out" if fam == "equity_decumulator" else "up_and_out",
               "daily_shares": ds, "step_up_daily_shares": stds, "gearing_ratio": gear,
               "max_nominal_shares": max_nom, "guaranteed_period_end": protected_end,
               "settlement_style": "physical_spot", "net_premium": 0.0,
               "trade_date": trade_date, "expiry_date": final_date, "tenor_days": tenor},
    )


# ── 条款抽取：MLI Booster / Leverage Call Spread + KI Put ───────────────

def extract_mli_terms(pdf_source, pdf_password: str = "") -> DerivativeTerm:
    text = _read_pdf_text(pdf_source, pages=8, pdf_password=pdf_password)
    symbol = _f(r"Underlying Share.*?Bloomberg.*?:\s*([A-Z0-9]{1,6}\s+[A-Z]{2})", text, flags=re.I | re.S) or ""
    symbol = symbol.split()[0]
    ccy = _f(r"([A-Z]{3})-Denominated", text) or "USD"
    # 4-month / 12-month in title
    months = _ff(r"A\s+(\d+)-month", text) or 4.0
    tenor = int(months * 30)

    # 先抓表格块：Underlying Share / Initial / KI / Strike 四列下方连续三行 USD 数值（真实样本 132/133/134）
    initial = ki_price = strike_price = 0.0
    mtab = re.search(
        r"Underlying Share \(Bloomberg Ticker\).*?Initial Price.*?Knock-in Price.*?Strike Price,? K.*?"
        r"[A-Za-z .()]+\([A-Z0-9]{1,6}\s+[A-Z]{2}\).*?USD\s*([\d,]+\.\d+).*?USD\s*([\d,]+\.\d+).*?USD\s*([\d,]+\.\d+)",
        text, re.I | re.S)
    if mtab:
        initial = float(mtab.group(1).replace(",", ""))
        ki_price = float(mtab.group(2).replace(",", ""))
        strike_price = float(mtab.group(3).replace(",", ""))
    else:
        initial = _ff(r"Initial Price\s*:?\s*USD\s*([\d,]+\.\d+)", text) or 0.0
        ki_price = _ff(r"Knock-in Price\s*:?\s*USD\s*([\d,]+\.\d+)", text) or 0.0
        strike_price = _ff(r"Strike Price.*?USD\s*([\d,]+\.\d+)", text) or 0.0

    ki_pct = _ff(r"Knock-in Price\s*\((\d+(?:\.\d+)?)% of Initial Price\)", text) or ((ki_price / initial) * 100 if initial else 0.0)
    strike_pct = _ff(r"Strike Price,?\s*K\s*\((\d+(?:\.\d+)?)% of Initial Price\)", text) or ((strike_price / initial) * 100 if initial else 100.0)
    max_up = _ff(r"maximum return of\s*(\d+(?:\.\d+)?)%", text) or _ff(r"Maximum Appreciation.*?(\d+(?:\.\d+)?)%", text) or 50.0
    pf = _ff(r"Participation Factor \(PF\)\s*:?\s*(\d+(?:\.\d+)?)%", text) or 100.0
    return DerivativeTerm(
        product_family="equity_mli_booster",
        underlying_symbol=symbol,
        currency=ccy,
        tenor_days=tenor,
        source_file=str(pdf_source),
        terms={
            "initial_price": initial,
            "knock_in_price": ki_price,
            "strike_price": strike_price,
            "participation_factor": pf / 100.0,
            "max_upside_pct": max_up / 100.0,
            "strike_pct_initial": strike_pct / 100.0,
            "knock_in_pct_initial": ki_pct / 100.0,
            "knock_in_direction": "down_and_in",
            "settlement_style": "physical",
            "principal_protected_if_no_ki": True,
            "tenor_days": tenor,
        },
    )


# ── 条款抽取：野村 irf「投资产品−理由记录」精简交易记录 ──────────────────
# irf-*.pdf 是野村每日下发的一页式成交记录（非完整条款单）：只给合约关键位
# (strike/KO)、名义/数量、成交/到期日、Gearing，无当期 MTM、无 DS/参考现价。
# → 薄记录：market_value_usd 留空(结构性产品栏不伪造 0)，payoff_* 全量重估待
#   完整 termsheet 导入再启用；总权益仍由月结单权威锚定，不依赖此记录逐日精算。

_IRF_MONTHS = {"JAN": 1, "FEB": 2, "MAR": 3, "APR": 4, "MAY": 5, "JUN": 6,
               "JUL": 7, "AUG": 8, "SEP": 9, "OCT": 10, "NOV": 11, "DEC": 12}


def _irf_trade_date(text: str) -> str:
    """`07−JUL−2026`（分隔符可能是 U+2212 减号或连字符）→ ISO `2026-07-07`。"""
    m = re.search(r"Trade Date:\s*(\d{1,2})[−\-]([A-Za-z]{3})[−\-](\d{4})", text)
    if not m:
        return ""
    mon = _IRF_MONTHS.get(m.group(2).upper(), 0)
    return f"{int(m.group(3)):04d}-{mon:02d}-{int(m.group(1)):02d}" if mon else ""


def _irf_maturity(ddmmyy: str) -> str:
    """产品串尾 `060727`（DDMMYY，20xx）→ ISO `2027-07-06`。"""
    return f"20{ddmmyy[4:6]}-{ddmmyy[2:4]}-{ddmmyy[0:2]}"


def _iso_days(a_iso: str, b_iso: str) -> int:
    from datetime import date
    try:
        a = date(*map(int, a_iso.split("-")))
        b = date(*map(int, b_iso.split("-")))
        return (b - a).days
    except (ValueError, TypeError):
        return 365


def _parse_irf_record(text: str, pdf_source) -> DerivativeTerm:
    """野村 irf 精简记录抽取。版式：
    `Investment Product: <类型 标的 交易所 | 标的 交易所 币种 类型> (strike/KO) DDMMYY`
    支持 累购/累沽/FCN 三族；lot_key=strike:DDMMYY 判别同标的多笔头寸。"""
    m = re.search(r"Investment Product:\s*(.+?)\(([\d.]+)\s*/\s*([\d.]+)\)\s*(\d{6})", text, re.S)
    if not m:
        raise ValueError("irf_product_not_found")
    desc = " ".join(m.group(1).split())
    up = desc.upper()
    strike, knock_out, ddmmyy = float(m.group(2)), float(m.group(3)), m.group(4)

    ccy = "USD"  # 野村日累购/减持均为 USD 计价；FCN 从产品串显式币种覆盖
    if "DECUMULATOR" in up:
        fam, ko_dir = "equity_decumulator", "down_and_out"
    elif "ACCUMULATOR" in up:
        fam, ko_dir = "equity_accumulator", "up_and_out"
    elif "FIXED COUPON NOTE" in up:
        fam, ko_dir = "equity_fcn", ""
    else:
        raise ValueError(f"irf_unknown_family:{desc[:40]}")

    if fam == "equity_fcn":
        fm = re.match(r"([A-Z0-9]{1,6})\s+[A-Z]{2}\s+([A-Z]{3})\s+FIXED COUPON NOTE", up)
        symbol = fm.group(1) if fm else up.split()[0]
        if fm:
            ccy = fm.group(2)
    else:
        am = re.search(r"(?:ACCUMULATOR|DECUMULATOR)\s+([A-Z0-9]{1,6})\s+[A-Z]{2}", up)
        symbol = am.group(1) if am else ""

    trade_date = _irf_trade_date(text)
    maturity = _irf_maturity(ddmmyy)
    tenor = _iso_days(trade_date, maturity)
    gearing = _ff(r"Gearing Ratio / Gearing Ratio Ideal Threshold:\s*([\d.]+)", text) or 0.0
    qm = re.search(r"Quantity / Notional:\s*(?:([A-Z]{3})\s+)?([\d,]+)", text)
    qty = float(qm.group(2).replace(",", "")) if qm else 0.0
    if qm and qm.group(1):
        ccy = qm.group(1)

    terms = {
        "strike": strike, "knock_out_price": knock_out, "knock_out_direction": ko_dir,
        "gearing_ratio": gearing, "trade_date": trade_date, "expiry_date": maturity,
        "maturity": maturity, "tenor_days": tenor, "market_value_usd": None,
        "settlement_style": "physical_spot", "source_kind": "nomura_irf",
    }
    if fam == "equity_fcn":
        terms["notional"] = qty
        terms["notional_ccy"] = ccy
    else:
        terms["afp"] = strike            # 累购/减持 AFP=行权价
        terms["max_nominal_shares"] = qty
    term = DerivativeTerm(product_family=fam, underlying_symbol=symbol, currency=ccy,
                          tenor_days=tenor, terms=terms, source_file=str(pdf_source))
    term.lot_key = f"{strike}:{ddmmyy}"
    return term


def extract_fcn_terms(pdf_source, pdf_password: str = "") -> DerivativeTerm:
    """FCN(固定息票票据) 抽取，三种版式同族 equity_fcn：
    野村 irf 精简记录 / 花旗完整条款单(Fixed Coupon Autocall Notes) / 巴克莱完整条款单(Daily Callable)。
    薄记录以条款价位为准，市值待重估。"""
    text = _read_pdf_text(pdf_source, pages=8, pdf_password=pdf_password)
    if "RATIONALE RECORD" in text:
        return _parse_irf_record(text, pdf_source)
    if "Fixed Coupon Autocall Notes" in text or "Autocall Barrier Level" in text:
        return _parse_citi_fcn(text, pdf_source)
    if "Daily Callable Fixed Coupon" in text or "Aggregate Nominal Amount" in text:
        return _parse_barclays_fcn(text, pdf_source)
    raise ValueError("fcn_full_termsheet_not_modeled")


def _fcn_thin_term(pdf_source, *, symbol, ccy, initial, strike, autocall, denom, notional,
                   trade_txt, mat_txt, source_kind, extra=None) -> DerivativeTerm:
    """完整 FCN 条款单 → equity_fcn 薄记录的公共装配（花旗/巴克莱共用）。
    lot_key=strike:DDMMYY 与野村 irf 一致 → 同产品重导入在持仓视图折叠为最新一期。"""
    def _iso(s):
        try:
            return datetime.strptime(s, "%d %B %Y").strftime("%Y-%m-%d")
        except ValueError:
            return ""
    trade_iso, mat_iso = _iso(trade_txt), _iso(mat_txt)
    tenor = _days_between(trade_txt, mat_txt) if trade_txt and mat_txt else 365
    ddmmyy = datetime.strptime(mat_txt, "%d %B %Y").strftime("%d%m%y") if mat_iso else "000000"
    terms = {
        "initial_price": initial, "strike": strike, "afp": strike,
        "knock_out_price": autocall, "autocall_barrier": autocall, "knock_out_direction": "up_and_out",
        "denomination": denom, "notional": notional, "notional_ccy": ccy,
        "trade_date": trade_iso, "expiry_date": mat_iso, "maturity": mat_iso,
        "tenor_days": tenor, "market_value_usd": None,
        "settlement_style": "cash_or_physical", "source_kind": source_kind,
    }
    if extra:
        terms.update(extra)
    term = DerivativeTerm(product_family="equity_fcn", underlying_symbol=symbol, currency=ccy,
                          tenor_days=tenor, terms=terms, source_file=str(pdf_source))
    term.lot_key = f"{strike}:{ddmmyy}"
    return term


def _parse_citi_fcn(text: str, pdf_source) -> DerivativeTerm:
    """花旗完整 FCN 条款单(Fixed Coupon Autocall Notes) → equity_fcn 薄记录。
    表格里三个价位(Initial/Strike/Autocall Barrier)标签与标的描述分块排列，故按
    "Barrier Level 标签之后首现的三个 4 位小数 USD 值" 顺序取值(Initial/Strike/Autocall)。"""
    symbol = (_f(r"([A-Z]{1,6})\s+[A-Z]{2}\s+Equity\b", text, flags=0) or "").strip()
    ccy = _f(r"Denomination\s+([A-Z]{3})\s+[\d,]", text) or "USD"
    denom = _ff(r"Denomination\s+[A-Z]{3}\s+([\d,]+(?:\.\d+)?)", text) or 0.0
    issue_size = _ff(r"Issue Size\s+[A-Z]{3}\s+([\d,]+(?:\.\d+)?)", text) or 0.0
    # 4 位小数 USD 值仅出现在标的价位区(息票 USD 333.35/面额 USD 50,000 均 ≤2 位)
    tail = text.split("Barrier Level", 1)[-1]
    lv = re.findall(r"USD\s*([\d,]+\.\d{3,})", tail)
    def _num(i):
        return float(lv[i].replace(",", "")) if i < len(lv) else 0.0
    initial, strike, autocall = _num(0), _num(1), _num(2)
    coupon_pa = _ff(r"approximately\s*([\d.]+)%\s*per annum", text) or 0.0
    return _fcn_thin_term(
        pdf_source, symbol=symbol, ccy=ccy, initial=initial, strike=strike, autocall=autocall,
        denom=denom, notional=issue_size,
        trade_txt=_f(r"Strike Date / Trade Date\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})", text) or "",
        mat_txt=_f(r"Maturity Date\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})", text) or "",
        source_kind="citi_fcn_termsheet", extra={"coupon_pa_pct": coupon_pa})


def _parse_barclays_fcn(text: str, pdf_source) -> DerivativeTerm:
    """巴克莱完整 FCN 条款单(Daily Callable Fixed Coupon, 经野村分销) → equity_fcn 薄记录。
    价位以具名定义给出：`"Initial Price" means USD X` / `Strike Price...being USD Y` / `Trigger Price...being USD Z`。
    coupon_pct 存每期票息原值(文档未给年化，不臆造 p.a.)。"""
    def _named_usd(label):
        m = re.search(rf'{label}[^"]*?\bUSD\s+([\d,]+\.\d+)', text, re.I | re.S)
        return float(m.group(1).replace(",", "")) if m else 0.0
    symbol = (_f(r"([A-Z]{1,6})\s+[A-Z]{2}\s+Equity\b", text, flags=0) or "").strip()
    ccy = _f(r"Specified Denomination\s+([A-Z]{3})\s+[\d,]", text) or "USD"
    denom = _ff(r"Specified Denomination\s+[A-Z]{3}\s+([\d,]+(?:\.\d+)?)", text) or 0.0
    notional = _ff(r"Aggregate Nominal Amount\s+[A-Z]{3}\s+([\d,]+(?:\.\d+)?)", text) or 0.0
    initial = _named_usd(r'"Initial Price"\s+means')
    strike = _named_usd(r'Strike Price"\s+means')
    autocall = _named_usd(r'Trigger Price"\s+means')
    coupon = _ff(r"Interest Rate\s+([\d.]+)%", text) or 0.0
    # 到期锚用 Redemption Date（scheduled to be …）；缺则退回 Final Valuation Date
    mat_txt = (_f(r"scheduled to be (\d{1,2}\s+[A-Za-z]+\s+\d{4})", text)
               or _f(r"Final Valuation Date\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})", text) or "")
    return _fcn_thin_term(
        pdf_source, symbol=symbol, ccy=ccy, initial=initial, strike=strike, autocall=autocall,
        denom=denom, notional=notional,
        trade_txt=_f(r"Trade Date\s+(\d{1,2}\s+[A-Za-z]+\s+\d{4})", text) or "",
        mat_txt=mat_txt, source_kind="barclays_fcn_termsheet", extra={"coupon_pct": coupon})


# ── FCN 条款单：Word(.docx) 版式（招银国际分销，摩根士丹利发行 Worst-of Autocall）──────────

def _docx_tables(raw: bytes) -> list[list[list[str]]]:
    """读 .docx 全部表格 → 行×格文本；合并单元格产生的"连续相等"文本折叠为一格。
    ponytail: 折叠连续相等即可去除合并列冗余(招银 basket 表每列跨 6~8 格)；相邻两列文本恰好全等的概率
    在价位/标签场景可忽略(strike=74.43%×initial 必不等于 initial)。要精确 grid 再上 gridSpan 解析。"""
    import io
    from docx import Document
    tables: list[list[list[str]]] = []
    for t in Document(io.BytesIO(raw)).tables:
        rows: list[list[str]] = []
        for r in t.rows:
            cells: list[str] = []
            for c in r.cells:
                txt = (c.text or "").replace("\xa0", " ").strip()
                if not cells or cells[-1] != txt:
                    cells.append(txt)
            rows.append(cells)
        tables.append(rows)
    return tables


def _first_date_str(s: str) -> str:
    m = re.search(r"(\d{1,2}\s+[A-Za-z]+\s+\d{4})", s or "")
    return m.group(1) if m else ""


def _dt_iso(s: str) -> str:
    d = _first_date_str(s)
    for f in ("%d %B %Y", "%d %b %Y"):
        try:
            return datetime.strptime(d, f).strftime("%Y-%m-%d")
        except ValueError:
            continue
    return ""


def _usd_num(s: str) -> float:
    m = re.search(r"[\d,]+\.?\d*", s or "")
    return float(m.group(0).replace(",", "")) if m else 0.0


def extract_cmbi_fcn_docx(raw: bytes) -> DerivativeTerm:
    """招银国际分销的 Word 条款单(Final Termsheet) → equity_fcn 记录。
    支持 Worst-of 多标的：underlying_symbol 用 'NVDA/SOXX' 篮子标签，逐腿明细存 terms['underlyings']；
    标量价位(initial/strike/knock_out)取首腿供 barrier/展示读取(条款期各腿同 %，实际最劣按收盘定)。"""
    tables = _docx_tables(raw)
    kv: dict[str, str] = {}
    for tbl in tables:
        for row in tbl:
            if len(row) == 2 and row[0] and row[0] != row[1]:
                kv.setdefault(row[0], row[1])

    def kget(*needles: str) -> str:
        for k, v in kv.items():
            kl = k.lower()
            if all(nd.lower() in kl for nd in needles):
                return v
        return ""

    def pct(s: str) -> float | None:
        m = re.search(r"([\d.]+)\s*%", s or "")
        return float(m.group(1)) if m else None

    ccy = (kget("Settlement Currency") or "USD").strip()[:3] or "USD"
    notional = _usd_num(kget("Aggregate Notional"))
    strike_pct = pct(kget("Strike Price"))          # 74.43
    autocall_pct = pct(kget("Autocallable Price"))  # 120
    coupon_pct = pct(kget("Coupon"))                # 1 (每观察期)
    trade_txt = _first_date_str(kget("Trade Date"))
    mat_txt = _first_date_str(kget("Maturity Date"))
    isin_raw = kget("ISIN")
    isin, _, common = isin_raw.partition("/")
    isin, common = isin.strip(), common.strip()

    # 标的篮子：定位含 INITIAL PRICE + BLOOMBERG 的表头行，按表头名映射列，读数字行(首格为序号)
    basket: list[dict] = []
    for tbl in tables:
        hidx = next((i for i, row in enumerate(tbl)
                     if "INITIAL PRICE" in " ".join(row).upper()
                     and "BLOOMBERG" in " ".join(row).upper()), None)
        if hidx is None:
            continue
        hdr = [c.upper() for c in tbl[hidx]]

        def col(*names: str) -> int:
            return next((j for j, h in enumerate(hdr) if any(n in h for n in names)), -1)

        ci = {k: col(*n) for k, n in {
            "name": ("UNDERLYING SEC",), "bbg": ("BLOOMBERG",), "init": ("INITIAL",),
            "strike": ("STRIKE",), "auto": ("AUTOCALL",), "exch": ("EXCHANGE",)}.items()}
        for row in tbl[hidx + 1:]:
            if not row or not row[0].strip().isdigit():
                continue
            def cell(key: str) -> str:
                j = ci[key]
                return row[j] if 0 <= j < len(row) else ""
            bbg = cell("bbg")
            sym = (bbg.split()[0] if bbg else cell("name").split()[0] if cell("name") else "SP").upper()
            basket.append({
                "symbol": sym, "name": cell("name"), "bloomberg": bbg,
                "initial_price": _usd_num(cell("init")), "strike_price": _usd_num(cell("strike")),
                "autocall_price": _usd_num(cell("auto")), "exchange": cell("exch")})
        break

    und_label = "/".join(b["symbol"] for b in basket) or "SP"
    mat_iso = _dt_iso(mat_txt)
    tenor = _days_between(trade_txt, mat_txt) if trade_txt and mat_txt else 0
    terms: dict = {
        "product_type": "autocallable_worst_of_fcn", "worst_of": len(basket) > 1,
        "underlyings": basket,
        "strike_pct": (strike_pct / 100.0) if strike_pct is not None else None,
        "autocall_pct": (autocall_pct / 100.0) if autocall_pct is not None else None,
        "coupon_pct": coupon_pct, "coupon_frequency": "per_observation_period",
        "notional": notional, "denomination": notional, "notional_ccy": ccy, "issue_price_pct": 100.0,
        "trade_date": _dt_iso(trade_txt), "strike_date": _dt_iso(kget("Strike Date")),
        "issue_date": _dt_iso(kget("Issue Date")), "final_valuation_date": _dt_iso(kget("Final Valuation Date")),
        "expiry_date": mat_iso, "maturity": mat_iso, "tenor_days": tenor,
        "autocall_observation_start": _dt_iso(kget("Autocallable Observation Period")),
        "autocall_frequency": "daily", "knock_out_direction": "up_and_out",
        "settlement_style": "cash_or_physical", "market_value_usd": None,
        "issuer": kget("Issuer"), "guarantor": kget("Guarantor"), "distributor": "cmbi",
        "isin": isin, "common_code": common, "source_kind": "cmbi_termsheet_docx",
    }
    if basket:  # 标量首腿：供 barrier_status/all-accounts 展示读取(worst-of 逐腿真值在 underlyings)
        b0 = basket[0]
        terms.update({"initial_price": b0["initial_price"], "strike": b0["strike_price"],
                      "afp": b0["strike_price"], "knock_out_price": b0["autocall_price"],
                      "autocall_barrier": b0["autocall_price"]})
    term = DerivativeTerm(product_family="equity_fcn", underlying_symbol=und_label,
                          currency=ccy, tenor_days=tenor, terms=terms)
    term.lot_key = f"{isin}:{mat_iso}" if isin else f"{und_label}:{mat_iso}"
    return term


def _cmbi_docx_selfcheck() -> None:
    import io
    try:
        from docx import Document
    except ImportError:
        return  # python-docx 未装(生产已装 1.2.0)：跳过自检
    econ = [("Issuer", "Morgan Stanley B.V. (not rated)"), ("Guarantor", "Morgan Stanley (A-)"),
            ("Trade Date", "26 May 2026"), ("Strike Date", "26 May 2026"), ("Issue Date", "2 June 2026"),
            ("Final Valuation Date", "2 September 2026"),
            ("Maturity Date", "4 September 2026, subject to adjustment in accordance with..."),
            ("Aggregate Notional Amount of the Notes", 'United States Dollar 1,060,000 ("USD")'),
            ("Specified Denomination", 'USD 1,060,000 per Note ("Par")'), ("Settlement Currency", "USD"),
            ("Strike Price", "equal to 74.43% of the Initial Price of such Underlying Security."),
            ("Autocallable Price", "equal to 120% of the Initial Price of such Underlying Security."),
            ("ISIN/Common Code", "XS3372957897/337295789"),
            ("Autocallable Observation Period", "From and including 2 July 2026 to the Final Valuation Date."),
            ("Coupon Amount", '"Coupon" means 1% per Observation Period.')]
    d = Document()
    t1 = d.add_table(rows=len(econ), cols=2)
    for i, (k, v) in enumerate(econ):
        t1.rows[i].cells[0].text = k
        t1.rows[i].cells[1].text = v
    bk_rows = [
        ["k", "UNDERLYING SECURITY", "BLOOMBERG CODE", "INITIAL PRICE",
         "STRIKE PRICE", "AUTOCALLABLE PRICE", "EXCHANGE"],
        ["1", "NVIDIA CORP", "NVDA UQ Equity", "USD 216.5100",
         "USD 161.1484", "USD 259.8120", "Nasdaq - All Markets"],
        ["2", "ISHARES SEMICONDUCTOR ETF", "SOXX UQ Equity", "USD 557.4000",
         "USD 414.8728", "USD 668.8800", "Nasdaq - All Markets"]]
    t2 = d.add_table(rows=len(bk_rows), cols=7)
    for i, r in enumerate(bk_rows):
        for j, val in enumerate(r):
            t2.rows[i].cells[j].text = val
    buf = io.BytesIO()
    d.save(buf)
    term = extract_cmbi_fcn_docx(buf.getvalue())
    assert term.product_family == "equity_fcn", term.product_family
    assert term.underlying_symbol == "NVDA/SOXX", term.underlying_symbol
    u = term.terms["underlyings"]
    assert len(u) == 2 and u[0]["symbol"] == "NVDA" and u[1]["symbol"] == "SOXX", u
    assert abs(u[0]["initial_price"] - 216.51) < 1e-6 and abs(u[1]["strike_price"] - 414.8728) < 1e-6, u
    assert abs(term.terms["strike_pct"] - 0.7443) < 1e-9, term.terms["strike_pct"]
    assert abs(term.terms["autocall_pct"] - 1.20) < 1e-9, term.terms["autocall_pct"]
    assert abs(term.terms["coupon_pct"] - 1.0) < 1e-9, term.terms["coupon_pct"]
    assert term.terms["maturity"] == "2026-09-04", term.terms["maturity"]
    assert term.terms["isin"] == "XS3372957897" and term.terms["common_code"] == "337295789", term.terms
    assert term.lot_key == "XS3372957897:2026-09-04", term.lot_key
    assert term.tenor_days == 101, term.tenor_days
    print("cmbi docx termsheet selfcheck 通过")


# ── 场景收益引擎 ─────────────────────────────────────────────────────────

def payoff_accumulator(term: DerivativeTerm, final_price: float, *,
                       knock_out_happened: bool = False, days_observed: int | None = None) -> dict:
    """简化场景引擎：按终值相对 AFP/KO 估算累计/减持股数与盈亏（静态近似，供报告风险提示）。

    - Accumulator：终值 < AFP → step-up 累积更多股（更危险）
    - Decumulator：终值 > AFP → step-up 减持更多股（上涨踏空风险）
    """
    t = term.terms
    ds = t.get("daily_shares", 0)
    stds = t.get("step_up_daily_shares", 0)
    afp = t.get("afp", 0.0)
    days = term.tenor_days if days_observed is None else int(days_observed)

    if term.product_family == "equity_decumulator":
        if knock_out_happened:
            shares = ds  # 敲出即终止：按 1 天 DS 保守近似（路径依赖不假装精确）
        else:
            shares = days * (ds if final_price <= afp else stds)
        proceeds = shares * afp
        market_value = shares * final_price
        return {"shares_decumulated": shares, "proceeds": proceeds, "market_value": market_value,
                "pnl": proceeds - market_value}

    # accumulator
    if knock_out_happened:
        shares = ds  # 敲出即终止：按 1 天 DS 保守近似（路径依赖不假装精确）
    else:
        shares = days * (ds if final_price >= afp else stds)
    cost = shares * afp
    mtm = shares * final_price
    return {"shares_acquired": shares, "cost": cost, "market_value": mtm, "pnl": mtm - cost}


def payoff_mli_booster(term: DerivativeTerm, final_price: float, *,
                       knock_in_happened: bool = False, investment_amount: float = 1.0) -> dict:
    """MLI Booster 到期收益：3 段式。返回到期兑付金额（以 investment_amount=1 归一化）与收益率。"""
    t = term.terms
    S0 = t.get("initial_price", 0.0)
    if S0 <= 0:
        return {"redemption": investment_amount, "return_pct": 0.0}
    strike = S0 * t.get("strike_pct_initial", 1.0)
    ki = S0 * t.get("knock_in_pct_initial", 0.0)
    pf = t.get("participation_factor", 1.0)
    cap = t.get("max_upside_pct", 0.5)
    upside = max(final_price / strike - 1.0, 0.0)
    upside = min(upside * pf, cap)
    if not knock_in_happened or final_price >= strike:
        redemption = investment_amount * (1.0 + upside)
    else:
        # down-and-in put：跌破 strike 后按标的跌幅承损
        redemption = investment_amount * max(final_price / strike, 0.0)
    return {"redemption": redemption, "return_pct": (redemption / investment_amount - 1.0) * 100,
            "knock_in_price": ki, "strike_price": strike}


def _is_indicative_intro(text: str) -> bool:
    """产品介绍/推介材料（indicative term sheet），非已成交持仓 → 不计入持仓。

    判别信号（真实花旗/野村样本全库无误判）：出现 "Indicative Terms"（推介用假设条款），
    且既无 "Final Terms"（已确认条款单）、也无野村 irf "RATIONALE RECORD"（成交薄记录）。
    共有的免责声明（"discussion purposes only"/"neither an offer"）是两类共用样板，不作判据。
    """
    if not re.search(r"Indicative Terms", text, re.I):
        return False
    return not re.search(r"Final Terms", text, re.I) and "RATIONALE RECORD" not in text


def classify_pdf(pdf_source, pdf_password: str = "") -> str:
    """日常文件快速分类：fund_report / accumulator / decumulator / mli / fcn / product_intro / other。

    产品介绍/推介材料（indicative term sheet，无成交/持仓）判为 product_intro，不落衍生品持仓——
    否则花旗 Step up AQ/DQ 等推介稿会被当作累购/累沽持仓列入持仓列表（用户反馈的根因）。
    """
    text = _read_pdf_text(pdf_source, pages=2, pdf_password=pdf_password)
    if _is_indicative_intro(text):
        return "product_intro"
    if "Master Fund Highlights" in text or "Financial Statements" in text:
        return "fund_report"
    # 完整条款单标题含 "Daily ... Accumulator/Decumulator"；野村 irf 精简记录只在产品串里点名
    # "EQUITY ACCUMULATOR/COVERED DECUMULATOR"（大写）→ 二者同判家族，抽取端自适应两种版式。
    low = text.lower()
    if re.search(r"daily(?: securities)? decumulator|covered decumulator|equity decumulator", low):
        return "decumulator"
    if re.search(r"daily(?: securities)? accumulator|equity accumulator", low):
        return "accumulator"
    # FCN(固定息票票据/每日可赎回)不是 MLI Booster；单列 fcn，避免 reextract 误走 extract_mli_terms 贴错家族。
    if "fixed coupon note" in low or "daily callable fixed coupon" in low \
            or "fixed coupon autocall" in low:  # 花旗完整条款单标题
        return "fcn"
    if "Market Linked Instrument" in text or "Leverage Call Spread" in text:
        return "mli"
    return "other"


def validate_derivative_term(term: DerivativeTerm, *, today: str = "") -> str:
    """导入前合理性护栏：返回非空原因串 = 不该入库（疑似解析残缺/错位）；空串 = 通过。

    衍生品/结构性产品被有意跳过股票的「陈旧/骤降」两道快照护栏（账户级头寸与快照时效无关，见
    importer._import_statement 注释），故此处补一道针对**条款数值本身**的合理性校验——否则解析
    错位（小数点/日期/版式列偏移）产出的荒谬记录会静默入库并展示给用户，误导持仓与风险判断。

    只查**能明确判定为错**的信号，宁松勿误拒（私行大额头寸不设名义/市值上限）：
      1. product_family / underlying_symbol 缺失 → 空壳记录（如风险披露页被误当条款单，实测有此样本）。
      2. notional 非正或非数 → 名义金额按定义恒正，<=0 是抽取错位。
      3. maturity 落在荒谬窗口（早于今年前 5 年 / 晚于今年后 30 年）→ 日期/数字错位
         （如把 strike 74.43 当年份、或版式偏移抽错列）。近期已到期票是正常历史，读时另有到期剔除。
    # ponytail: 不校 strike_pct/coupon_pct 等百分比——各字段单位口径不统一（fraction vs 1.0=1%），
    #   区间护栏易误拒；拿到"百分号错位"真实样本再按字段单独加。
    """
    from bottleneck_hunter.watchlist.store_base import _today
    if not (term.product_family or "").strip():
        return "missing_product_family"
    if not (term.underlying_symbol or "").strip():
        return "missing_underlying_symbol"
    t = term.terms or {}
    notional = t.get("notional")
    if notional is not None:
        try:
            if float(notional) <= 0:
                return f"nonpositive_notional:{notional}"
        except (TypeError, ValueError):
            return f"bad_notional:{notional!r}"
    maturity = (t.get("maturity") or t.get("expiry_date") or "").strip()
    m = re.match(r"^(\d{4})-", maturity)
    if m:                                    # 只在能抽出年份时判窗；非 ISO 版式不因格式误拒
        y = int(m.group(1))
        cur_year = int((today or _today())[:4])
        if not (cur_year - 5 <= y <= cur_year + 30):
            return f"maturity_out_of_window:{maturity}"
    return ""


def save_derivative_term(wl_store, term: DerivativeTerm, *, source_file_name: str, source_file_hash: str,
                         broker: str, rationale_ref: str = "", account_ref: str = "", lot_key: str = "") -> str:
    import json
    import uuid
    account_ref = wl_store.resolve_vip_account_ref(account_ref) if hasattr(wl_store, "resolve_vip_account_ref") else (account_ref or "").strip()
    lot_key = (lot_key or "").strip()  # 同标的多笔头寸判别键；条款单单条路径留空(行为不变)
    # 幂等：重复上传同一文件保留原 id/created_at，但刷新 terms_json/currency——否则 parser 升级
    # (如野村到期日改 ISO)的修复永远进不了旧行。与 vip_imports「重导刷新既有行」同理(见
    # importer.create_vip_import 注释：is_redo 直接 return 会令代码升级后的字段永远写不进旧行)。
    conn = wl_store._connect()
    try:
        q, p = wl_store._filtered(
            "SELECT id FROM vip_derivative_terms WHERE account_ref=? AND source_file_hash=? AND product_family=? AND underlying_symbol=? AND lot_key=?",
            (account_ref, source_file_hash, term.product_family, term.underlying_symbol, lot_key))
        row = conn.execute(q, p).fetchone()
    finally:
        conn.close()
    if row:
        with wl_store._write_conn() as conn:
            q2, p2 = wl_store._filtered(
                "UPDATE vip_derivative_terms SET terms_json=?, currency=? WHERE id=?",
                (json.dumps(term.terms, ensure_ascii=False), term.currency, row["id"]))
            conn.execute(q2, p2)
        return row["id"]
    did = uuid.uuid4().hex[:12]
    with wl_store._write_conn() as conn:
        conn.execute(
            f"""INSERT INTO vip_derivative_terms
               (id, source_file_name, source_file_hash, broker, product_family, underlying_symbol,
                currency, terms_json, rationale_ref, account_ref, lot_key, created_at{wl_store._user_insert_cols()}{wl_store._market_insert_cols()})
               VALUES (?,?,?,?,?,?,?,?,?,?,?,?{wl_store._user_insert_vals()}{wl_store._market_insert_vals()})""",
            (did, source_file_name, source_file_hash, broker, term.product_family, term.underlying_symbol,
             term.currency, json.dumps(term.terms, ensure_ascii=False), rationale_ref, account_ref, lot_key, datetime.now().isoformat())
            + wl_store._user_insert_params() + wl_store._market_insert_params(),
        )
    return did


def update_derivative_term(wl_store, did: str, term: DerivativeTerm) -> bool:
    """按 id 覆盖某条衍生品条款的可变字段（重抽回填用）。用户/市场隔离由 _filtered 保证。

    返回是否命中并更新（未命中该用户名下的 id 则 False）。
    """
    import json
    with wl_store._write_conn() as conn:
        q, p = wl_store._filtered(
            """UPDATE vip_derivative_terms
               SET product_family=?, underlying_symbol=?, currency=?, terms_json=?
               WHERE id=?""",
            (term.product_family, term.underlying_symbol, term.currency,
             json.dumps(term.terms, ensure_ascii=False), did),
        )
        cur = conn.execute(q, p)
        return cur.rowcount > 0


def delete_derivative_term(wl_store, did: str) -> bool:
    """按 id 删除一条衍生品条款记录（误导入清理用，如推介稿/风险披露稿被误判为持仓）。

    用户/市场隔离由 _filtered 保证；无补删机制此前是历史脏数据只能永久残留的根因（用户反馈）。
    返回是否命中并删除（未命中该用户名下的 id 则 False）。
    """
    with wl_store._write_conn() as conn:
        q, p = wl_store._filtered("DELETE FROM vip_derivative_terms WHERE id=?", (did,))
        cur = conn.execute(q, p)
        return cur.rowcount > 0


def list_derivative_terms(wl_store, limit: int = 50, account_ref: str = "") -> list[DerivativeTerm]:
    import json
    account_ref = wl_store.resolve_vip_account_ref(account_ref) if hasattr(wl_store, "resolve_vip_account_ref") else (account_ref or "").strip()
    conn = wl_store._connect()
    try:
        q, p = wl_store._filtered(
            "SELECT * FROM vip_derivative_terms WHERE account_ref=? ORDER BY created_at DESC LIMIT ?", (account_ref, limit))
        rows = [dict(r) for r in conn.execute(q, p).fetchall()]
    finally:
        conn.close()
    out = []
    for r in rows:
        t = json.loads(r["terms_json"] or "{}")
        out.append(DerivativeTerm(product_family=r["product_family"], underlying_symbol=r["underlying_symbol"],
                                  currency=r["currency"], tenor_days=int(t.get("tenor_days", 0) or 0),
                                  terms=t, source_file=r["source_file_name"], id=r["id"],
                                  lot_key=r.get("lot_key", "") or ""))
    return out


def _merge_fcn_terms(rows: list[dict]) -> dict:
    """同一笔 FCN 的多份记录(条款单 docx + 月结单)合并为一条 terms。
    条款结构字段(strike_pct/autocall/coupon/underlyings/价位…)以 docx 条款单为准(最全、无 MTM)；
    MTM/名义/当期结余以最新月结单为准。都缺则各自兜底。rows 已按 (created_at,id) 升序。"""
    docx = next((r for r in reversed(rows) if r["_terms"].get("source_kind") == "cmbi_termsheet_docx"), None)
    latest_stmt = next((r for r in reversed(rows) if r["_terms"].get("source_kind") != "cmbi_termsheet_docx"), None)
    base = dict((docx or rows[-1])["_terms"])  # 条款基底：优先 docx，否则最新那份
    if latest_stmt is not None:  # 月结单独有的鲜活值覆盖上来（条款单 market_value_usd 恒 None）
        st = latest_stmt["_terms"]
        for k in ("market_value_usd", "market_value_nominal", "notional", "underlying_name", "nominal_ccy"):
            v = st.get(k)
            if v is not None:
                base[k] = v
    return base


def list_derivative_terms_all_accounts(wl_store, limit: int = 200, account_ref: str | None = None) -> list[dict]:
    """结构性产品/衍生品"当前"明细（account_ref=None 全账户总览；给定则单账户「持仓 Tab」，两处共此一份口径）。

    同一 (account_ref, family, lot_key) 折成一条——lot_key(=ISIN:到期)已唯一标识一笔，**不含 underlying_symbol**：
    否则条款单 docx(篮子标签 NVDA/SOXX) 与月结单(NVDA) 同一笔 FCN 折叠键不匹配 → 双计虚增(commit 本轮根因)。
    多份记录跨源合并 terms（条款以 docx 为准、MTM 以最新月结单为准，见 _merge_fcn_terms），而非整份取最新丢字段。
    再剔除已过到期日(北京日期口径)的头寸——到期的旧票(如招银 CMBIGP step-up notes)不该再挂在当前持仓里。
    """
    import json
    from collections import defaultdict

    from bottleneck_hunter.watchlist.store_base import _today

    conn = wl_store._connect()
    try:
        base = ("SELECT account_ref, product_family, underlying_symbol, currency, terms_json, lot_key, "
                "source_file_name, created_at, id FROM vip_derivative_terms")
        args: tuple = ()
        if account_ref is not None:
            ref = wl_store.resolve_vip_account_ref(account_ref) if hasattr(wl_store, "resolve_vip_account_ref") else (account_ref or "").strip()
            base += " WHERE account_ref=?"
            args = (ref,)
        q, p = wl_store._filtered(base + " ORDER BY created_at ASC, id ASC", args, table="vip_derivative_terms")
        rows = [dict(r) for r in conn.execute(q, p).fetchall()]
    finally:
        conn.close()
    for r in rows:
        r["_terms"] = json.loads(r["terms_json"] or "{}")
    groups: dict = defaultdict(list)
    for r in rows:  # 折叠键去 symbol：同 (account, family, lot_key) 即同一笔（lot_key=ISIN:到期 已唯一）。
        # 但 lot_key 空(条款单版式没抽到 strike/ISIN)时**不折叠**——退回按 id 各自独立：否则同账户同 family
        # 不同标的的多份无 key 条款单(如花旗 MU 与 9988 accumulator)会塌进同一空键被错折、丢数据。
        lk = (r.get("lot_key") or "").strip()
        key = (r.get("account_ref") or "", r["product_family"], lk or f"\x00id={r['id']}")
        groups[key].append(r)
    today = _today()
    items = []
    for grp in groups.values():
        t = _merge_fcn_terms(grp)
        maturity = t.get("maturity") or t.get("expiry_date") or ""
        if maturity and maturity < today:  # 已过到期日 → 不再是当前持仓（ISO 日期串直接比较）
            continue
        latest = grp[-1]  # 展示元数据(symbol/币种/来源文件)取最新一份；symbol 篮子标签优先(更全)
        symbol = next((r["underlying_symbol"] for r in reversed(grp) if "/" in (r["underlying_symbol"] or "")),
                      latest["underlying_symbol"])
        items.append({
            "id": latest["id"],  # 最新一份的 id：reextract 覆盖最新单，符合「跟最新单」语义
            "product_family": latest["product_family"],
            "underlying_symbol": symbol,
            "currency": latest["currency"],
            "tenor_days": int(t.get("tenor_days", 0) or 0),
            "market_value_usd": t.get("market_value_usd"),
            "notional": t.get("notional"),
            "maturity": maturity,
            "terms": t,
            "source_file": latest["source_file_name"],
            "account_ref": latest.get("account_ref") or "",
        })
    items.sort(key=lambda x: x["maturity"] or "", reverse=True)
    return items[:limit]


def demo() -> None:
    # 教科书基准
    p = bs_price(100, 100, 1, 0.05, 0.2, True)
    assert abs(p - 10.4506) < 1e-3, p
    # IV 往返
    iv = implied_vol(p, 100, 100, 1, 0.05, True)
    assert iv and abs(iv - 0.2) < 1e-3, iv
    # Accumulator 场景
    acc = DerivativeTerm("equity_accumulator", "MU", "USD", 365, {"afp": 100.0, "daily_shares": 3, "step_up_daily_shares": 6})
    r = payoff_accumulator(acc, 80.0, knock_out_happened=False, days_observed=10)
    assert r["shares_acquired"] == 60 and r["pnl"] < 0
    # Booster 场景
    mli = DerivativeTerm("equity_mli_booster", "MU", "USD", 120,
                         {"initial_price": 100.0, "participation_factor": 1.0,
                          "max_upside_pct": 0.5, "strike_pct_initial": 1.0, "knock_in_pct_initial": 0.5379})
    assert payoff_mli_booster(mli, 130.0, knock_in_happened=False)["return_pct"] > 0
    assert payoff_mli_booster(mli, 80.0, knock_in_happened=True)["return_pct"] < 0
    _lot_key_selfcheck()
    _sanity_guard_selfcheck()
    _delete_selfcheck()
    _all_accounts_selfcheck()
    _irf_selfcheck()
    _citi_fcn_selfcheck()
    _barclays_fcn_selfcheck()
    _cmbi_docx_selfcheck()
    _intro_guard_selfcheck()
    print("derivatives demo 通过")


def _sanity_guard_selfcheck() -> None:
    """导入护栏 validate_derivative_term：坏数据被拦、正常大额头寸放行、边界日期不误伤。"""
    good = DerivativeTerm("equity_fcn", "NVDA", "USD", 100,
                          {"notional": 1060000, "maturity": "2026-09-04", "strike_pct": 0.7443})
    assert validate_derivative_term(good, today="2026-08-04") == "", "正常 FCN 被误拒"
    # 私行大额名义(千万级)不设上限 → 放行
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "TSLA", "USD", 90, {"notional": 50_000_000, "maturity": "2027-01-01"}),
        today="2026-08-04") == ""
    # 空壳：无 family / 无标的（风险披露页被误当条款单）
    assert validate_derivative_term(DerivativeTerm("", "NVDA", "USD", 1, {})).startswith("missing_product_family")
    assert validate_derivative_term(DerivativeTerm("equity_fcn", "", "USD", 1, {})).startswith("missing_underlying")
    # 名义非正 / 非数（抽取错位）
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"notional": 0})).startswith("nonpositive_notional")
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"notional": -5})).startswith("nonpositive_notional")
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"notional": "abc"})).startswith("bad_notional")
    # 到期日荒谬（strike 74 被当年份 / 版式偏移）→ 拦；近期已到期(去年)属正常历史 → 放行
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"maturity": "0074-01-01"}),
        today="2026-08-04").startswith("maturity_out_of_window")
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"maturity": "2099-01-01"}),
        today="2026-08-04").startswith("maturity_out_of_window")
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"maturity": "2025-06-30"}),
        today="2026-08-04") == "", "近期已到期票不该被合理性护栏拦(读时另有到期剔除)"
    # 非 ISO 日期版式不因格式误拒（无年份可判 → 跳过窗判）
    assert validate_derivative_term(
        DerivativeTerm("equity_fcn", "NVDA", "USD", 1, {"maturity": "04SEP26"}),
        today="2026-08-04") == ""
    print("sanity_guard 自检通过")


def _intro_guard_selfcheck() -> None:
    """产品介绍判别自检：Indicative Terms 且无 Final Terms/RATIONALE RECORD → intro；
    已成交单(Final Terms)与野村 irf(RATIONALE RECORD)不误判。共有免责声明不作判据。"""
    disclaimer = "This document is neither an offer to sell; for discussion purposes only."
    intro = "A 1-Year Daily Securities Accumulator\nIndicative Terms as of 02 June 2025\n" + disclaimer
    assert _is_indicative_intro(intro), "推介稿(Indicative Terms)未判为 intro"
    final = "Final Terms\nA 1-Year Daily Securities Accumulator\n" + disclaimer
    assert not _is_indicative_intro(final), "已成交单(Final Terms)被误判为 intro"
    irf = "INVESTMENT PRODUCT − RATIONALE RECORD\nIndicative Terms\n" + disclaimer
    assert not _is_indicative_intro(irf), "野村 irf(RATIONALE RECORD)被误判为 intro"
    assert not _is_indicative_intro(disclaimer), "仅免责声明不应判为 intro"


def _irf_selfcheck() -> None:
    """野村 irf 精简记录抽取自检（内联伪造文本，不碰真实 PII PDF）：
    三族家族/标的/strike/KO/币种/到期/lot_key 正确；同标的不同 strike → lot_key 不折叠。"""
    acc = ("INVESTMENT PRODUCT − RATIONALE RECORD\nTrade Date: \n07−JUL−2026\n"
           "Account Number: \n22704339\nInvestment Product: \nOTC EQUITY ACCUMULATOR BE UN \n"
           "(169.803/278.765) 060727\nQuantity / Notional: \n1500 Shares\n"
           "Gearing Ratio / Gearing Ratio Ideal Threshold: \n6.01 / 3\n")
    t = _parse_irf_record(acc, "irf-a.pdf")
    assert t.product_family == "equity_accumulator" and t.underlying_symbol == "BE", t
    assert t.terms["strike"] == 169.803 and t.terms["knock_out_price"] == 278.765
    assert t.terms["knock_out_direction"] == "up_and_out" and t.currency == "USD"
    assert t.terms["trade_date"] == "2026-07-07" and t.terms["maturity"] == "2027-07-06"
    assert t.terms["max_nominal_shares"] == 1500 and t.terms["gearing_ratio"] == 6.01
    assert t.terms["market_value_usd"] is None and t.lot_key == "169.803:060727"

    dec = acc.replace("OTC EQUITY ACCUMULATOR BE UN", "OTC EQUITY COVERED DECUMULATOR PLTR UW") \
             .replace("(169.803/278.765)", "(178.1962/128.5101)")
    d = _parse_irf_record(dec, "irf-d.pdf")
    assert d.product_family == "equity_decumulator" and d.underlying_symbol == "PLTR", d
    assert d.terms["knock_out_direction"] == "down_and_out"

    fcn = ("INVESTMENT PRODUCT − RATIONALE RECORD\nTrade Date: \n13−JUL−2026\n"
           "Investment Product: \nNVDA US USD FIXED COUPON NOTE \n(182.9692/200.7864) 141226\n"
           "Quantity / Notional: \nUSD 1000000\n"
           "Gearing Ratio / Gearing Ratio Ideal Threshold: \n6.02 / 3\n")
    f = _parse_irf_record(fcn, "irf-f.pdf")
    assert f.product_family == "equity_fcn" and f.underlying_symbol == "NVDA", f
    assert f.currency == "USD" and f.terms["notional"] == 1000000
    assert f.terms["maturity"] == "2026-12-14"

    # 同标的不同 strike → lot_key 不折叠（野村双 BE accumulator 场景）
    acc2 = acc.replace("(169.803/278.765) 060727", "(155.500/260.000) 260727")
    t2 = _parse_irf_record(acc2, "irf-a2.pdf")
    assert t2.lot_key != t.lot_key, "同标的不同 strike 的 lot_key 折叠了"


def _all_accounts_selfcheck() -> None:
    """全账户当前明细(list_derivative_terms_all_accounts)：同一笔多份结单折一条 + 已过到期剔除 + MTM/名义透出。
    这正是招银 05/06/07 三份月结单致 FCN 三倍虚增、及到期 CMBIGP 仍挂列表的两处根因。"""
    import tempfile
    from pathlib import Path

    from bottleneck_hunter.watchlist.store import WatchlistStore
    with tempfile.TemporaryDirectory() as d:
        wl = WatchlistStore(db_path=Path(d) / "t.db").for_user("u1").for_market("us_stock")
        fut = DerivativeTerm("equity_fcn", "NVDA", "USD", 120,
                             {"market_value_usd": 12345.6, "notional": 1000000, "maturity": "2099-12-31"})
        past = DerivativeTerm("equity_fcn", "CMBIGP", "USD", 60, {"market_value_usd": 999.0, "maturity": "2000-01-01"})
        common = dict(broker="cmbi", account_ref="acc")
        # 同一笔 NVDA FCN 两份月结单(不同 file_hash，同 lot_key) → 应折成一条
        save_derivative_term(wl, fut, source_file_name="m05.pdf", source_file_hash="h05", lot_key="L1", **common)
        save_derivative_term(wl, fut, source_file_name="m06.pdf", source_file_hash="h06", lot_key="L1", **common)
        save_derivative_term(wl, past, source_file_name="m05.pdf", source_file_hash="h05b", lot_key="L2", **common)
        items = list_derivative_terms_all_accounts(wl)
        syms = [i["underlying_symbol"] for i in items]
        assert syms.count("NVDA") == 1, f"同笔 FCN 未折叠：{syms}"
        assert "CMBIGP" not in syms, f"已到期未剔除：{syms}"
        nv = next(i for i in items if i["underlying_symbol"] == "NVDA")
        assert nv["market_value_usd"] == 12345.6 and nv["notional"] == 1000000, nv
        assert nv["maturity"] == "2099-12-31" and nv["account_ref"], nv

        # 跨源同 ISIN：条款单 docx(篮子标签 NVDA/SOXX) + 月结单(NVDA) 同 lot_key → 折一条 + terms 合并（本轮根因）
        docx_t = DerivativeTerm("equity_fcn", "NVDA/SOXX", "USD", 101,
                                {"maturity": "2099-06-30", "strike_pct": 0.7443, "autocall_pct": 1.20,
                                 "coupon_pct": 1.0, "underlyings": [{"symbol": "NVDA"}, {"symbol": "SOXX"}],
                                 "market_value_usd": None, "source_kind": "cmbi_termsheet_docx"})
        stmt_t = DerivativeTerm("equity_fcn", "NVDA", "USD", 101,
                                {"maturity": "2099-06-30", "market_value_usd": 55555.5, "notional": 1060000,
                                 "underlying_name": "NVIDIA CORP"})
        save_derivative_term(wl, stmt_t, source_file_name="m07.pdf", source_file_hash="hs7",
                             lot_key="XS337:2099-06-30", **common)
        save_derivative_term(wl, docx_t, source_file_name="ts.docx", source_file_hash="hd",
                             lot_key="XS337:2099-06-30", **common)
        it2 = list_derivative_terms_all_accounts(wl)
        wo = [i for i in it2 if i["maturity"] == "2099-06-30"]
        assert len(wo) == 1, f"跨源同 ISIN 未折叠：{[i['underlying_symbol'] for i in it2]}"
        m = wo[0]
        assert m["underlying_symbol"] == "NVDA/SOXX", f"篮子标签未保留：{m['underlying_symbol']}"
        assert m["terms"]["strike_pct"] == 0.7443 and m["terms"]["coupon_pct"] == 1.0, "条款(docx)未合入"
        assert m["market_value_usd"] == 55555.5 and m["notional"] == 1060000, "MTM/名义(月结单)未合入"

        # 空 lot_key 不折叠：同账户同 family 不同标的的无 key 条款单(花旗 MU 与 9988)必须各自独立，不被错折
        for sym in ("MU", "BABA"):
            nokey = DerivativeTerm("equity_accumulator", sym, "USD", 365, {"maturity": "2099-01-01"})
            save_derivative_term(wl, nokey, source_file_name=f"{sym}.pdf", source_file_hash=f"h{sym}",
                                 lot_key="", **common)
        it3 = list_derivative_terms_all_accounts(wl)
        acc_syms = [i["underlying_symbol"] for i in it3 if i["product_family"] == "equity_accumulator"]
        assert "MU" in acc_syms and "BABA" in acc_syms, f"空 lot_key 跨标的被错折：{acc_syms}"


def _delete_selfcheck() -> None:
    """delete_derivative_term 自检：命中删除返回 True 且真实清除；未命中(id 不存在)返回 False。
    历史误导入(推介稿/风险披露稿被误判为持仓)无补删机制是脏数据永久残留的根因(用户反馈)。"""
    import tempfile
    from pathlib import Path

    from bottleneck_hunter.watchlist.store import WatchlistStore
    with tempfile.TemporaryDirectory() as d:
        wl = WatchlistStore(db_path=Path(d) / "t.db").for_user("u1").for_market("us_stock")
        junk = DerivativeTerm("equity_accumulator", "9988", "USD", 0, {})
        did = save_derivative_term(wl, junk, source_file_name="junk.pdf", source_file_hash="hj",
                                   broker="citi", account_ref="acc")
        assert delete_derivative_term(wl, did), "命中的 id 删除应返回 True"
        assert list_derivative_terms(wl, account_ref="acc") == [], "删除后应不再出现在列表里"
        assert not delete_derivative_term(wl, did), "已删除的 id 再删应返回 False（非报错）"
        assert not delete_derivative_term(wl, "nonexistent"), "不存在的 id 应返回 False"


def _lot_key_selfcheck() -> None:
    """lot_key 去重键自检(临时 DB)：同 (file,family,symbol) 不同 lot_key → 两条；同 lot_key → 幂等一条。
    这正是野村双 ORCL accumulator 此前被静默折叠的根因。"""
    import tempfile
    from pathlib import Path

    from bottleneck_hunter.watchlist.store import WatchlistStore
    with tempfile.TemporaryDirectory() as d:
        wl = WatchlistStore(db_path=Path(d) / "t.db").for_user("u1").for_market("us_stock")
        t1 = DerivativeTerm("equity_accumulator", "ORCL", "USD", 100, {"market_value_usd": -3245.18})
        t2 = DerivativeTerm("equity_accumulator", "ORCL", "USD", 110, {"market_value_usd": -1100.0})
        common = dict(source_file_name="s.pdf", source_file_hash="h", broker="nomura", account_ref="acc")
        id_a = save_derivative_term(wl, t1, lot_key="244.0263:161026", **common)
        id_b = save_derivative_term(wl, t2, lot_key="230.5649:261026", **common)
        id_a2 = save_derivative_term(wl, t1, lot_key="244.0263:161026", **common)  # 重放 → 幂等
        assert id_a != id_b, "不同 lot_key 折叠成一条（去重键回归）"
        assert id_a == id_a2, "同 lot_key 未幂等"
        assert len(list_derivative_terms(wl, account_ref="acc")) == 2, "双 ORCL 应两条"


def _citi_fcn_selfcheck() -> None:
    """花旗完整 FCN 条款单抽取自检（内联伪造版式，不碰真实 PII PDF）：
    标的/币种/三价位/到期/名义/lot_key 正确；lot_key 与野村 irf 同格式(strike:DDMMYY)可折叠。"""
    text = (
        "Fixed Coupon Autocall Notes Based Upon the Shares of NVIDIA Corporation\n"
        "Issue Size \nUSD 1,000,000 \nDenomination \nUSD 50,000 \n"
        "Strike Date / Trade Date \n27 May 2026 \nMaturity Date \n14 December 2026 \n"
        "Initial Level \nStrike Level \nAutocall \nBarrier Level \n"
        "NVDA UW \nEquity \nShare \nNASDAQ \n"
        "USD 211.3541 \nUSD 182.9692 \nUSD 200.7864 \n"
        "Strike Level \n86.57% of the Initial Level \n"
        "0.6667% per month (corresponding to approximately 8.00% per annum)\n"
    )
    t = _parse_citi_fcn(text, "termsheet.pdf")
    assert t.product_family == "equity_fcn" and t.underlying_symbol == "NVDA", t
    assert t.currency == "USD" and t.terms["notional"] == 1000000 and t.terms["denomination"] == 50000
    assert t.terms["initial_price"] == 211.3541 and t.terms["strike"] == 182.9692
    assert t.terms["autocall_barrier"] == 200.7864 and t.terms["coupon_pa_pct"] == 8.0
    assert t.terms["trade_date"] == "2026-05-27" and t.terms["maturity"] == "2026-12-14"
    assert t.terms["market_value_usd"] is None and t.lot_key == "182.9692:141226"


def _barclays_fcn_selfcheck() -> None:
    """巴克莱完整 FCN 条款单(Daily Callable, 具名价位) 抽取自检（内联伪造版式，不碰真实 PII PDF）。"""
    text = (
        "Daily Callable Fixed Coupon linked to NVIDIA CORP\n"
        "Trade Date \n13 July 2026 \nAggregate Nominal Amount \nUSD 1,000,000 \n"
        "Specified Denomination \nUSD 50,000 \nBloomberg Code (for identification purposes only) \n"
        "NVDA UW Equity \nInterest Rate \n0.6667% \n"
        '"Initial Price" means USD 206.1495, being the price on the Initial Valuation Date. \n'
        '"Trigger Price" means 95.00% of the Initial Price; being USD 195.8420 as at the Trade Date. \n'
        '"Strike Price" means 84.45% of the Initial Price; being USD 174.0933 as at the Trade Date. \n'
        "Two (2) Business Days immediately following the Final Valuation Date (scheduled to be 30 December 2026). \n"
    )
    t = _parse_barclays_fcn(text, "barclays.pdf")
    assert t.product_family == "equity_fcn" and t.underlying_symbol == "NVDA", t
    assert t.currency == "USD" and t.terms["notional"] == 1000000 and t.terms["denomination"] == 50000
    assert t.terms["initial_price"] == 206.1495 and t.terms["strike"] == 174.0933
    assert t.terms["autocall_barrier"] == 195.842 and t.terms["coupon_pct"] == 0.6667
    assert t.terms["trade_date"] == "2026-07-13" and t.terms["maturity"] == "2026-12-30"
    assert t.terms["market_value_usd"] is None and t.lot_key == "174.0933:301226"


if __name__ == "__main__":
    demo()
